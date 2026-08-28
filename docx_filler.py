"""Вставка тела ответа в файл-шаблон .docx.

Логика: открыть шаблон объекта, найти вводный абзац (тот, что начинается со слов
вида «ООО … является Подрядчиком/исполнителем …»), и вставить абзацы тела
СРАЗУ ПОСЛЕ него, скопировав его оформление (шрифт, отступы, выравнивание).
Всё остальное в бланке остаётся нетронутым.
"""

import copy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


class IntroParagraphNotFound(Exception):
    """Не удалось найти вводный абзац в шаблоне."""


def _find_intro_paragraph(doc: Document, intro_marker: str) -> Paragraph:
    marker = intro_marker.lower()
    for par in doc.paragraphs:
        if marker in par.text.lower():
            return par
    raise IntroParagraphNotFound(
        f"В шаблоне не найден вводный абзац с фразой «{intro_marker}»."
    )


def _insert_paragraph_after(ref_par: Paragraph, text: str, style_par: Paragraph) -> Paragraph:
    """Создать новый абзац сразу после ref_par, скопировав оформление из style_par."""
    new_p = OxmlElement("w:p")

    # Копируем свойства абзаца (отступы, выравнивание, интервалы).
    style_pPr = style_par._p.find(qn("w:pPr"))
    if style_pPr is not None:
        new_p.append(copy.deepcopy(style_pPr))

    ref_par._p.addnext(new_p)
    new_par = Paragraph(new_p, ref_par._parent)

    run = new_par.add_run(text)
    # Копируем оформление шрифта из первого прогона вводного абзаца.
    if style_par.runs:
        style_rPr = style_par.runs[0]._r.find(qn("w:rPr"))
        if style_rPr is not None:
            run._r.insert(0, copy.deepcopy(style_rPr))

    return new_par


def fill_template(
    template_path: Path,
    body_paragraphs: list[str],
    output_path: Path,
    intro_marker: str,
) -> Path:
    """Вставить абзацы тела в шаблон и сохранить копию по output_path."""
    doc = Document(str(template_path))
    intro = _find_intro_paragraph(doc, intro_marker)

    # Вставляем абзацы по очереди, каждый — после предыдущего, оформление берём с вводного.
    anchor = intro
    for text in body_paragraphs:
        anchor = _insert_paragraph_after(anchor, text, style_par=intro)

    doc.save(str(output_path))
    return output_path
